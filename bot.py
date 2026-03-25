import asyncio
import json
import logging
import os
from dataclasses import dataclass
import io
from typing import Optional

from aiogram import Bot, Dispatcher, F
from aiogram.enums import ParseMode
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    CallbackQuery,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    KeyboardButton,
    Message,
    ReplyKeyboardRemove,
    ReplyKeyboardMarkup,
)
from aiogram.types.input_file import BufferedInputFile
from dotenv import load_dotenv
from openai import AsyncOpenAI
from pypdf import PdfReader
from docx import Document as DocxDocument


logging.basicConfig(level=logging.INFO)
log = logging.getLogger("gpt_consultant_bot")

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
STORE_PATH = os.path.join(PROJECT_DIR, "gpt_store.json")


DEFAULT_SYSTEM_PROMPT = (
    "Ты — деловой консультант по подготовке документации и выдаче таблиц соответствия.\n"
    "Отвечай коротко, по делу, без токсичности.\n\n"
    "ЖЁСТКОЕ ТЗ (строгий сценарий):\n"
    "1) Сначала запроси аукционные документы.\n"
    "2) После того как пользователь прислал аукционные документы — запроси приложение.\n"
    "3) После приложения — запроси руководство по эксплуатации.\n"
    "4) После получения всех трёх типов материалов выдай таблицу соответствия в двух версиях:\n"
    "   - Версия WORD: таблица в формате, пригодном для вставки в Word (например, Markdown/HTML-таблица).\n"
    "   - Версия PDF: таблица в формате, пригодном для экспорта/печати (например, тот же HTML/таблица).\n\n"
    "Правила:\n"
    "- Не используй вопросительные предложения и знак `?`.\n"
    "- Проси документы/материалы императивно: «Пришли…», «Нужен…».\n"
    "- Если данных пока недостаточно — запрашивай следующий документ по списку.\n"
    "- Для таблицы соответствия используй извлечённые из присланных материалов сущности/разделы.\n"
)


@dataclass(frozen=True)
class Settings:
    bot_token: str
    openai_api_key: str
    openai_model: str
    admin_telegram_ids: list[int]
    gpt_system_prompt: str


def _split_ints(value: str) -> list[int]:
    parts = [p.strip() for p in (value or "").split(",")]
    return [int(p) for p in parts if p]


def load_settings() -> Settings:
    load_dotenv()
    bot_token = os.getenv("BOT_TOKEN", "").strip()
    openai_api_key = os.getenv("OPENAI_API_KEY", "").strip()
    openai_model = os.getenv("OPENAI_MODEL", "gpt-5.4").strip()
    admin_ids_raw = os.getenv("ADMIN_TELEGRAM_IDS", "").strip()
    admin_telegram_ids = _split_ints(admin_ids_raw)
    gpt_system_prompt = os.getenv("GPT_SYSTEM_PROMPT", "").strip()

    missing: list[str] = []
    if not bot_token:
        missing.append("BOT_TOKEN")
    if not openai_api_key:
        missing.append("OPENAI_API_KEY")
    if not admin_telegram_ids:
        missing.append("ADMIN_TELEGRAM_IDS")

    if missing:
        raise RuntimeError(f"Missing env vars: {', '.join(missing)}")

    return Settings(
        bot_token=bot_token,
        openai_api_key=openai_api_key,
        openai_model=openai_model,
        admin_telegram_ids=admin_telegram_ids,
        gpt_system_prompt=gpt_system_prompt,
    )


def ensure_store() -> dict:
    if not os.path.exists(STORE_PATH):
        initial_prompt = os.getenv("GPT_SYSTEM_PROMPT", "").strip() or DEFAULT_SYSTEM_PROMPT
        initial_model = os.getenv("OPENAI_MODEL", "").strip() or "gpt-5.4"
        with open(STORE_PATH, "w", encoding="utf-8") as f:
            json.dump(
                {
                    "version": 1,
                    "gpt_system_prompt": initial_prompt,
                    "openai_model": initial_model,
                },
                f,
                ensure_ascii=False,
                indent=2,
            )

    with open(STORE_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def normalize_system_prompt(prompt: str) -> str:
    """
    If an old store prompt still contains the former persona,
    replace it with current DEFAULT_SYSTEM_PROMPT to match new requirements.
    """
    p = (prompt or "").lower()
    if "продавщиц" in p or "хамоват" in p:
        return DEFAULT_SYSTEM_PROMPT
    return prompt or DEFAULT_SYSTEM_PROMPT


def save_store(data: dict) -> None:
    with open(STORE_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def get_admin_kb(*, is_admin: bool) -> ReplyKeyboardMarkup | ReplyKeyboardRemove:
    if not is_admin:
        return ReplyKeyboardRemove()
    return ReplyKeyboardMarkup(keyboard=[[KeyboardButton(text="АДМИН")]], resize_keyboard=True)


def admin_inline_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="Показать ТЗ", callback_data="admin:show_prompt")],
            [InlineKeyboardButton(text="Изменить ТЗ", callback_data="admin:edit_prompt")],
            [InlineKeyboardButton(text="Изменить модель", callback_data="admin:edit_model")],
        ]
    )


def chunk_text(text: str, limit: int = 3800) -> list[str]:
    if not text:
        return [""]
    if len(text) <= limit:
        return [text]
    return [text[i : i + limit] for i in range(0, len(text), limit)]


class AdminStates(StatesGroup):
    WAIT_PROMPT = State()
    WAIT_MODEL = State()


MAX_HISTORY_TURNS = 10
# Store full extracted text; we will chunk when sending to GPT.
MAX_DOC_CHARS = 2_000_000
ALLOWED_DOC_EXTS = {".pdf", ".docx", ".doc"}
MAX_TABLE_CHARS = 2_000_000

FIXED_COLUMNS = [
    "№ п/п",
    "Наименование параметра, соответствующего заявке на закупку",
    "Соответствует / Не соответствует",
    "Ссылка на документ (раздел/пункт/страница)",
]


def _get_ext(file_name: str) -> str:
    _, ext = os.path.splitext((file_name or "").lower())
    return ext


def _truncate(text: str, max_chars: int) -> str:
    text = text or ""
    if len(text) <= max_chars:
        return text
    return text[:max_chars]


async def download_document_bytes(bot: Bot, file_id: str) -> bytes:
    tg_file = await bot.get_file(file_id)
    buf = io.BytesIO()
    # aiogram v3: download_file writes into a buffer-like destination
    await bot.download_file(tg_file.file_path, destination=buf)
    return buf.getvalue()


async def extract_pdf_text(pdf_bytes: bytes) -> str:
    def _extract() -> str:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        parts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            parts.append(page_text)
        return "\n".join(parts).strip()

    return await asyncio.to_thread(_extract)


async def extract_docx_text(docx_bytes: bytes) -> str:
    def _extract() -> str:
        doc = DocxDocument(io.BytesIO(docx_bytes))
        parts: list[str] = [p.text for p in doc.paragraphs if (p.text or "").strip()]
        return "\n".join(parts).strip()

    return await asyncio.to_thread(_extract)


async def extract_document_text(file_name: str, doc_bytes: bytes) -> str:
    ext = _get_ext(file_name)
    if ext == ".pdf":
        return await extract_pdf_text(doc_bytes)
    if ext in {".docx", ".doc"}:
        # Telegram may send .doc; parsing .doc as docx won't work reliably.
        # We'll accept only docx content; if it's a true .doc, extracted text may fail.
        return await extract_docx_text(doc_bytes)
    return ""

def _normalize_whitespace(text: str) -> str:
    return "\n".join([line.rstrip() for line in (text or "").splitlines()]).strip()


def make_docx_bytes(title: str, markdown_table: str) -> bytes:
    # Backward-compat wrapper (not used anymore)
    doc = DocxDocument()
    if title:
        doc.add_heading(title, level=1)
    doc.add_paragraph(markdown_table or "")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def normalize_table(table: dict) -> tuple[str, list[str], list[list[str]]]:
    title = str(table.get("title") or "Таблица соответствия")
    # Force exact order like the sample.
    columns = list(FIXED_COLUMNS)
    rows = table.get("rows")

    norm_rows: list[list[str]] = []
    if isinstance(rows, list):
        for r in rows:
            if not isinstance(r, list):
                continue
            rr = [str(x) for x in r]
            norm_rows.append(rr)

    # Fill numbering if missing/empty
    fixed = []
    for i, row in enumerate(norm_rows):
        row0 = row[0] if len(row) > 0 else ""
        if not row0.strip():
            if len(row) == 0:
                row = [str(i + 1)]
            else:
                row = [str(i + 1)] + row[1:]
        fixed.append(row)
    norm_rows = fixed

    # Pad/truncate cells to match columns
    out_rows: list[list[str]] = []
    for row in norm_rows:
        row = row[: len(columns)]
        while len(row) < len(columns):
            row.append("")
        out_rows.append([cell.replace("\n", " ").strip() for cell in row])

    return title, columns, out_rows


def make_docx_table_bytes(title: str, columns: list[str], rows: list[list[str]]) -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col

    for row in rows:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            cells[i].text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()



async def gpt_table_json(client: AsyncOpenAI, model: str, system_prompt: str, payload: str) -> dict:
    """
    Ask GPT to return a JSON object with:
    { "title": str, "columns": [str...], "rows": [[cell...], ...] }
    """
    instruction = (
        "Сформируй таблицу соответствия и верни строго JSON без Markdown и без пояснений.\n"
        "Колонки строго в этом порядке:\n"
        f'1) {FIXED_COLUMNS[0]}\n'
        f'2) {FIXED_COLUMNS[1]}\n'
        f'3) {FIXED_COLUMNS[2]}\n'
        f'4) {FIXED_COLUMNS[3]}\n'
        "Формат:\n"
        '{ "title": "...", "columns": ["№ п/п","...","...","..."], "rows": [["1","...","...","..."], ...] }\n'
        "Если не уверен — пиши «Неизвестно» и давай ссылку на документ/раздел где возможно указано."
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": instruction + "\n\nДанные:\n" + payload},
    ]
    resp = await client.chat.completions.create(model=model, messages=messages, temperature=0.2)
    content = (resp.choices[0].message.content or "").strip()
    return json.loads(content)


def chunk_for_gpt(text: str, limit: int = 14000) -> list[str]:
    text = text or ""
    if len(text) <= limit:
        return [text]
    return [text[i : i + limit] for i in range(0, len(text), limit)]


async def gpt_table_json_chunked(client: AsyncOpenAI, model: str, system_prompt: str, payload: str) -> dict:
    """
    Send ALL payload to GPT, split into parts if large, then ask for JSON table.
    """
    parts = chunk_for_gpt(payload, limit=14000)
    messages: list[dict[str, str]] = [{"role": "system", "content": system_prompt}]
    for idx, part in enumerate(parts, start=1):
        messages.append(
            {
                "role": "user",
                "content": f"Часть {idx}/{len(parts)} данных (сохрани в контексте):\n{part}",
            }
        )
    messages.append(
        {
            "role": "user",
            "content": (
                "Теперь сформируй таблицу соответствия и верни строго JSON.\n"
                "Колонки строго в этом порядке:\n"
                + "\n".join([f"{i+1}) {c}" for i, c in enumerate(FIXED_COLUMNS)])
                + "\nФормат:\n"
                '{ "title": "...", "columns": ["№ п/п","...","...","..."], "rows": [["1","...","...","..."], ...] }'
            ),
        }
    )
    resp = await client.chat.completions.create(model=model, messages=messages, temperature=0.2)
    content = (resp.choices[0].message.content or "").strip()
    return json.loads(content)


def json_table_to_markdown(table: dict) -> tuple[str, str]:
    title = str(table.get("title") or "Таблица соответствия")
    columns = table.get("columns") or []
    rows = table.get("rows") or []
    if not isinstance(columns, list) or not columns:
        columns = ["Раздел/требование", "Источник", "Комментарий"]
    # Ensure all rows are lists of correct width
    norm_rows: list[list[str]] = []
    for r in rows if isinstance(rows, list) else []:
        if not isinstance(r, list):
            continue
        rr = [str(x) for x in r][: len(columns)]
        while len(rr) < len(columns):
            rr.append("")
        norm_rows.append(rr)

    header = "| " + " | ".join(columns) + " |"
    sep = "| " + " | ".join(["---"] * len(columns)) + " |"
    body = "\n".join(["| " + " | ".join([cell.replace("\n", " ") for cell in row]) + " |" for row in norm_rows])
    md = "\n".join([header, sep, body]).strip()
    return title, md


def next_missing_field(data: dict) -> str:
    if not data.get("auction_text"):
        return "auction"
    if not data.get("appendix_text"):
        return "appendix"
    if not data.get("manual_text"):
        return "manual"
    return "done"


async def gpt_reply(client: AsyncOpenAI, model: str, system_prompt: str, user_text: str, history: list[dict]) -> str:
    # Simple chat history: last few turns.
    messages: list[dict[str, str]] = [{"role": "system", "content": system_prompt}]
    messages.extend(history[-10:])
    messages.append({"role": "user", "content": user_text})

    resp = await client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.7,
    )
    return (resp.choices[0].message.content or "").strip()


async def cmd_start(message: Message, state: FSMContext) -> None:
    await state.clear()
    is_admin = message.from_user.id in SETTINGS.admin_telegram_ids
    await message.answer(
        "Здравствуй.\n"
        "Бот поможет подготовить пакет материалов и выдать таблицу соответствия в форматах WORD и PDF.\n\n"
        "Сейчас шаг 1: пришли аукционные документы (текстом/содержимым/выжимкой). "
        "Дальше бот попросит приложение и руководство по эксплуатации, после чего сформирует таблицу.",
        reply_markup=get_admin_kb(is_admin=is_admin),
        parse_mode=ParseMode.MARKDOWN,
    )


async def on_admin_text(message: Message, state: FSMContext) -> None:
    if message.from_user.id not in SETTINGS.admin_telegram_ids:
        await message.answer("Нет доступа.")
        return
    await state.clear()
    await message.answer("Админ-панель:", reply_markup=admin_inline_kb())


async def on_admin_callback(callback: CallbackQuery, state: FSMContext) -> None:
    if callback.from_user.id not in SETTINGS.admin_telegram_ids:
        await callback.answer("Нет доступа.", show_alert=True)
        return

    store = ensure_store()
    data = callback.data or ""

    if data == "admin:show_prompt":
        prompt = str(store.get("gpt_system_prompt") or DEFAULT_SYSTEM_PROMPT)
        model = str(store.get("openai_model") or SETTINGS.openai_model)
        await callback.message.edit_text(
            f"Текущие значения:\nМодель: `{model}`\n\nТЗ отправлю целиком (частями, если нужно).",
            parse_mode=ParseMode.MARKDOWN,
        )
        for i, chunk in enumerate(chunk_text(prompt), start=1):
            await callback.message.answer(f"ТЗ (часть {i}):\n{chunk}")
        await callback.answer()
        return

    if data == "admin:edit_prompt":
        await state.set_state(AdminStates.WAIT_PROMPT)
        await callback.message.edit_text("Пришли новое ТЗ для GPT (одним сообщением).")
        await callback.answer()
        return

    if data == "admin:edit_model":
        await state.set_state(AdminStates.WAIT_MODEL)
        await callback.message.edit_text(
            "Пришли новое имя модели API (например, `gpt-5.4-mini`).",
            parse_mode=ParseMode.MARKDOWN,
        )
        await callback.answer()
        return

    await callback.answer()


async def on_admin_message(message: Message, state: FSMContext) -> None:
    if message.from_user.id not in SETTINGS.admin_telegram_ids:
        return

    st = await state.get_state()
    if st == AdminStates.WAIT_PROMPT.state:
        text = (message.text or "").strip()
        if not text:
            await message.answer("Пусто. Пришли ТЗ снова.")
            return
        store = ensure_store()
        store["gpt_system_prompt"] = text
        save_store(store)
        await state.clear()
        await message.answer("ТЗ обновлено.", reply_markup=admin_inline_kb())
        return

    if st == AdminStates.WAIT_MODEL.state:
        text = (message.text or "").strip()
        if not text:
            await message.answer("Пусто. Пришли модель снова.")
            return
        store = ensure_store()
        store["openai_model"] = text
        save_store(store)
        await state.clear()
        await message.answer("Модель обновлена.", reply_markup=admin_inline_kb())
        return


async def on_user_message(message: Message, state: FSMContext) -> None:
    if message.text is None:
        return

    if message.text.strip() == "АДМИН":
        return

    # If admin is editing, don't treat messages as user chat.
    st = await state.get_state()
    if st and st.startswith("AdminStates"):
        return

    store = ensure_store()
    model = str(store.get("openai_model") or SETTINGS.openai_model)
    system_prompt = normalize_system_prompt(str(store.get("gpt_system_prompt") or DEFAULT_SYSTEM_PROMPT))

    data = await state.get_data()
    history: list[dict] = list(data.get("history") or [])
    user_text = message.text.strip()

    if not user_text:
        return

    await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")

    try:
        reply = await gpt_reply(OPENAI, model, system_prompt, user_text, history)
    except Exception:
        log.exception("GPT user message error")
        await message.answer("Не получилось обработать запрос. Попробуй ещё раз или коротко переформулируй.")
        return

    if not reply:
        await message.answer("Пусто вышло. Попробуй ещё раз или перефразируй.")
        return

    await message.answer(reply)

    # Save history (for context)
    history.append({"role": "user", "content": user_text})
    history.append({"role": "assistant", "content": reply})
    await state.update_data(history=history[-20:])


async def on_user_document(message: Message, state: FSMContext) -> None:
    if not message.document:
        return

    store = ensure_store()
    model = str(store.get("openai_model") or SETTINGS.openai_model)
    system_prompt = normalize_system_prompt(str(store.get("gpt_system_prompt") or DEFAULT_SYSTEM_PROMPT))

    file_name = message.document.file_name or "document"
    ext = _get_ext(file_name)
    if ext not in ALLOWED_DOC_EXTS:
        await message.answer("Поддерживаю только `PDF` и `DOCX` (и иногда `DOC`, если Telegram отдаст его корректно).")
        return

    caption = (message.caption or "").strip()

    # Acknowledge immediately so user sees progress.
    await message.answer("Материал получил. Обрабатываю, подожди немного.")
    await message.bot.send_chat_action(chat_id=message.chat.id, action="typing")

    try:
        doc_bytes = await download_document_bytes(message.bot, message.document.file_id)
        extracted = await extract_document_text(file_name, doc_bytes)
    except Exception:
        log.exception("Document processing failed")
        await message.answer("Не смог обработать документ. Попробуй отправить в PDF/DOCX ещё раз.")
        return

    extracted = _truncate(_normalize_whitespace(extracted), MAX_DOC_CHARS)
    if not extracted:
        extracted = "(текст из документа не извлёкся; если документ-скан, пришли текстом или перешли в OCR)"

    data = await state.get_data()
    missing = next_missing_field(data)

    # Store documents by the strict scenario order.
    if missing == "auction":
        await state.update_data(auction_text=extracted)
        await message.answer("Принял аукционные документы. Теперь пришли приложение.")
        return
    if missing == "appendix":
        await state.update_data(appendix_text=extracted)
        await message.answer("Принял приложение. Теперь пришли руководство по эксплуатации.")
        return
    if missing == "manual":
        await state.update_data(manual_text=extracted)

    # If we got here, all 3 are present -> build the table and send as DOCX + PDF files.
    data = await state.get_data()
    auction_text = str(data.get("auction_text") or "")
    appendix_text = str(data.get("appendix_text") or "")
    manual_text = str(data.get("manual_text") or "")

    payload = (
        "Аукционные документы:\n" + auction_text + "\n\n"
        "Приложение:\n" + appendix_text + "\n\n"
        "Руководство по эксплуатации:\n" + manual_text + "\n\n"
        + (f"Комментарий пользователя к последнему файлу: {caption}\n" if caption else "")
    )

    try:
        await message.answer("Все материалы получены. Формирую таблицу соответствия. Подожди.")
        table = await gpt_table_json_chunked(OPENAI, model, system_prompt, payload)
    except Exception:
        log.exception("GPT failed for table generation")
        await message.answer("Не смог сформировать таблицу. Попробуй ещё раз или уменьши объём документов.")
        return

    title, columns, rows = normalize_table(table)
    docx_bytes = await asyncio.to_thread(make_docx_table_bytes, title, columns, rows)

    await message.answer_document(
        BufferedInputFile(docx_bytes, filename="table.docx"),
        caption="Таблица соответствия (WORD).",
    )


SETTINGS = load_settings()
STORE = ensure_store()

OPENAI = AsyncOpenAI(api_key=SETTINGS.openai_api_key)


async def main() -> None:
    bot = Bot(token=SETTINGS.bot_token)
    storage = MemoryStorage()
    dp = Dispatcher(storage=storage)

    dp.message.register(cmd_start, Command("start"))
    dp.message.register(on_admin_text, F.text == "АДМИН")

    dp.callback_query.register(on_admin_callback, F.data.startswith("admin:"))
    dp.message.register(on_admin_message, AdminStates.WAIT_PROMPT)
    dp.message.register(on_admin_message, AdminStates.WAIT_MODEL)

    dp.message.register(on_user_document, F.document)
    dp.message.register(on_user_message)

    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())

